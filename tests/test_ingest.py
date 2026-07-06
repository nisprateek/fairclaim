"""T&C ingestion guards: size caps, file types, empty input (no network)."""

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from fairclaim.backend.security.ingest import MAX_BYTES, IngestError, ingest_file, ingest_pasted_text


def test_empty_paste_rejected():
    with pytest.raises(IngestError):
        ingest_pasted_text("   ")


def test_paste_is_size_capped():
    assert len(ingest_pasted_text("x" * (MAX_BYTES + 100))) == MAX_BYTES


def test_txt_upload_decodes():
    assert ingest_file("terms.txt", b"No refunds on faulty goods.") == "No refunds on faulty goods."


def test_oversized_upload_rejected():
    with pytest.raises(IngestError):
        ingest_file("terms.txt", b"x" * (MAX_BYTES + 1))


def test_unsupported_extension_rejected():
    with pytest.raises(IngestError):
        ingest_file("terms.exe", b"whatever")


def test_extension_check_is_case_insensitive():
    assert ingest_file("TERMS.TXT", b"All sales final.") == "All sales final."


def test_docx_upload_extracts_paragraph_text():
    doc = DocxDocument()
    doc.add_paragraph("Clause 1: No refunds.")
    doc.add_paragraph("Clause 2: Sold as seen.")
    buffer = BytesIO()
    doc.save(buffer)

    text = ingest_file("terms.docx", buffer.getvalue())
    assert "Clause 1: No refunds." in text
    assert "Clause 2: Sold as seen." in text


def test_pdf_upload_with_no_extractable_text_returns_empty():
    # A blank/scanned page must come back as empty text, not crash.
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    assert ingest_file("terms.pdf", buffer.getvalue()) == ""


def test_txt_with_invalid_utf8_is_replaced_not_crashed():
    text = ingest_file("terms.txt", b"No refunds \xff\xfe on sale items.")
    assert "No refunds" in text and "on sale items." in text
